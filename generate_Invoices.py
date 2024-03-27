from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
import os
from datetime import datetime
from dateutil.relativedelta import relativedelta


files_list = []
today = datetime.now().strftime("%A %d %B %Y")
date = datetime.now().strftime("%d-%m-%Y")
last_month = datetime.now() - relativedelta(months=+1)
last_month = last_month.strftime("%A %d %B %Y")

def csvReader():
    with open('Invoices.csv') as csv_file:
        csv_reader = csv.reader(csv_file, delimiter=',')
        line_count = 0
        for row in csv_reader:
            file_name = row[1] + '_Invoice_' + date + '.docx'
            if line_count == 0:
                line_count += 1
            else:
                line_count += 1
                document = Document('template.docx')

                p_ref = document.add_paragraph(row[0])
                p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_date = document.add_paragraph(date)
                p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                document.add_paragraph('Name: ' + row[1])
                document.add_paragraph('Website: ' + row[2])
                document.add_paragraph('Email: ' + row[3])
                document.add_paragraph('Address: ' + row[4] + ", " + row[5] + ", " + row[6] + ", " + row[7])
                document.add_paragraph('Invoice Period: ' + last_month + ' - ' + today)

                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.add_break()
                run.add_break()

                if row[9] != f'0':
                    table = document.add_table(4, 2)
                    heading_cells = table.rows[0].cells
                    heading_cells[0].text = 'Description'
                    heading_cells[1].text = 'Amount'
                    table.cell(1, 0).text = 'Monthly Fee'
                    table.cell(1, 1).text = ''.join(['£', format(float(row[8]), ',.2f')])
                    table.cell(2, 0).text = 'Additional Fee'
                    table.cell(2, 1).text = ''.join(['£', format(float(row[9]), ',.2f')])
                    table.cell(3, 0).text = 'Total Due'
                    table.cell(3, 1).text = ''.join(['£', format(float(row[10]), ',.2f')])

                else:
                    table = document.add_table(3, 2)
                    heading_cells = table.rows[0].cells
                    heading_cells[0].text = 'Description'
                    heading_cells[1].text = 'Amount'
                    table.cell(1, 0).text = 'Monthly Fee'
                    table.cell(1, 1).text = ''.join(['£', format(float(row[8]), ',.2f')])
                    table.cell(2, 0).text = 'Total Due'
                    table.cell(2, 1).text = ''.join(['£', format(float(row[10]), ',.2f')])
                
                table.style = 'Light Shading Accent 1'

                document.save(file_name)

        print(f'Processed {line_count-1} Entries.')
        
csvReader()